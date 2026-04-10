from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from PIL import Image, ImageTk
import os
import shutil
import json
import copy

def iter_block_items(parent, document_ref=None):
    """문단과 표를 원래 순서대로 순회"""
    parent_element = parent.element.body if hasattr(parent, 'element') and hasattr(parent.element, 'body') else parent._element
    owner = document_ref or parent

    for child in parent_element.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, owner)
        elif child.tag == qn('w:tbl'):
            yield Table(child, owner)

def extract_block_content(container, area='body'):
    """본문/머릿글/바닥글의 문단과 표를 구조적으로 추출"""
    items = []

    for block in iter_block_items(container):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                items.append({
                    'type': 'text',
                    'text': text,
                    'style': getattr(block.style, 'name', 'Normal'),
                    'area': area
                })

                if area != 'body':
                    continue

            if area == 'body':
                for run in block.runs:
                    if run._element.xpath('.//a:blip'):
                        blip = run._element.xpath('.//a:blip')[0]
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId in container.part.rels:
                            rel = container.part.rels[rId]
                            image_part = rel.target_part
                            image_data = image_part.blob
                            content_type = image_part.content_type
                            ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                            temp_dir = os.path.join(os.path.dirname(container.part.package.part_related_by(rId).partname if False else ''), '.temp_images')
        elif isinstance(block, Table):
            table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            if table_data:
                items.append({
                    'type': 'table',
                    'data': table_data,
                    'area': area
                })

    return items

def add_bordered_table(target, data):
    """지정한 컨테이너에 실선 테두리 표 추가"""
    if not data:
        return None

    if hasattr(target, 'add_table') and hasattr(target, 'sections'):
        table = target.add_table(rows=len(data), cols=len(data[0]))
    else:
        try:
            table = target.add_table(rows=len(data), cols=len(data[0]), width=Inches(6.5))
        except TypeError:
            table = target.add_table(rows=len(data), cols=len(data[0]))

    table.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/></w:tcBorders>')
            tcPr.append(tcBorders)
    return table

def clear_story_container(container):
    """머릿글/바닥글의 기존 문단/표 제거"""
    for element in list(container._element):
        if element.tag in (qn('w:p'), qn('w:tbl')):
            container._element.remove(element)

    if not container.paragraphs:
        container.add_paragraph()

def safe_style_name(paragraph):
    """문단 스타일명 안전 조회 (일부 문서의 스타일 오류 회피)"""
    try:
        style = paragraph.style
        if style is not None and hasattr(style, 'name'):
            return style.name
    except Exception:
        pass
    return 'Normal'

def load_existing_doc(file_path):
    """Word 문서에서 모든 내용 추출 (문단 및 이미지 순서대로)"""
    doc = Document(file_path)
    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
    
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    all_content = []
    image_counter = 0

    try:
        header_items = []
        for block in iter_block_items(doc.sections[0].header):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    header_items.append({
                        'type': 'text',
                        'text': text,
                        'style': safe_style_name(block),
                        'area': 'header'
                    })
            elif isinstance(block, Table):
                table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if table_data:
                    header_items.append({
                        'type': 'table',
                        'data': table_data,
                        'area': 'header'
                    })
        all_content.extend(header_items)
    except:
        pass

    for element in doc.element.body:
        if element.tag == qn('w:p'):
            para = Paragraph(element, doc)
            if para.text.strip():
                all_content.append({
                    'type': 'text',
                    'text': para.text,
                    'style': safe_style_name(para),
                    'area': 'body'
                })

            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    blip = run._element.xpath('.//a:blip')[0]
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        image_part = rel.target_part
                        image_data = image_part.blob
                        content_type = image_part.content_type
                        ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                        image_file = os.path.join(temp_dir, f"image_{image_counter}{ext}")
                        with open(image_file, 'wb') as f:
                            f.write(image_data)
                        all_content.append({
                            'type': 'image',
                            'path': image_file,
                            'element': element,  # 원본 w:p 요소 (deepcopy 사용 시 관계 보존)
                            'area': 'body'
                        })
                        image_counter += 1
        elif element.tag == qn('w:tbl'):
            table = Table(element, doc)
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if table_data:
                all_content.append({
                    'type': 'table',
                    'data': table_data,
                    'element': element,  # 병합 셀·서식 보존을 위해 원본 XML 요소 저장
                    'area': 'body'
                })

    try:
        footer_items = []
        for block in iter_block_items(doc.sections[0].footer):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    footer_items.append({
                        'type': 'text',
                        'text': text,
                        'style': safe_style_name(block),
                        'area': 'footer'
                    })
            elif isinstance(block, Table):
                table_data = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                if table_data:
                    footer_items.append({
                        'type': 'table',
                        'data': table_data,
                        'area': 'footer'
                    })
        all_content.extend(footer_items)
    except:
        pass
    
    return all_content

def extract_images_from_doc(file_path):
    """Word 문서에서 이미지 추출"""
    doc = Document(file_path)
    image_paths = []
    temp_dir = os.path.join(os.path.dirname(file_path), ".temp_images")
    
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_data = image_part.blob
                content_type = image_part.content_type
                ext = '.jpg' if 'jpeg' in content_type else '.png' if 'png' in content_type else '.gif'
                image_file = os.path.join(temp_dir, f"image_{len(image_paths)}{ext}")
                with open(image_file, 'wb') as f:
                    f.write(image_data)
                image_paths.append(image_file)
    except Exception as e:
        print(f"이미지 추출 오류: {e}")
    
    return image_paths

def create_ndt_procedure_doc(paragraphs, image_paths=None):
    """문단 정보와 이미지로 새로운 문서 생성"""
    doc = Document()
    
    for para_info in paragraphs:
        text = para_info.get('text', '').strip()
        if not text:
            continue
            
        style = para_info.get('style', 'Normal')
        
        # Heading 스타일 처리
        if 'Heading' in style:
            try:
                level = int(style.split()[-1])
            except:
                level = 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text, style=style)
    
    # 이미지 추가
    if image_paths:
        doc.add_heading('첨부 사진', level=1)
        for img_path in image_paths:
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass
    
    return doc

class NDTProcedureApp:
    CONFIG_FILE = "app_config.json"
    
    def __init__(self, root):
        self.root = root
        self.root.title("비파괴 검사 절차서 생성 및 관리 시스템")
        self.root.resizable(True, True)
        
        # 저장된 창 크기 복원 또는 기본값 사용
        geometry = self.load_window_geometry()
        self.root.geometry(geometry)
        
        # 창 크기 변경 이벤트 바인드
        self.root.bind('<Configure>', self.on_window_configure)
        
        self.paragraphs = []
        self.image_paths = []
        self.content = []
        self.source_file = None  # 원본 Word 파일 경로 (바닥글/헤더 이미지 보존용)
        self.standards = {
            "ASME Section V, Article 4 (PAUT 기본 절차)":
                "ASME Section V, Article 4 - Ultrasonic Examination Methods (PAUT 기본 절차)\n\n적용 범위:\n위상배열 초음파검사(PAUT)의 핵심 절차 요구사항을 규정하는 기본 코드.\n\n주요 요구사항:\n- 위상배열 프로브 사양 및 선정 기준\n- 스캔 계획(Scan Plan) 수립 및 시뮬레이션\n- 보정 블록(Calibration Block) 규격 및 보정 절차\n- 감도 설정 및 DAC/TCG 적용\n- 결함 탐지, 위치, 크기 측정 기준\n- 기록 요구사항: A-scan, S-scan 데이터 보존\n\n검사원 자격: ASNT SNT-TC-1A 또는 CP-189 기준 Level II 이상\n보정 주기: 검사 전·후 및 8시간마다 보정 확인 필요",
            "ASME Section V, Article 4, Appendix III (PAUT 전용 요구사항)":
                "ASME Section V, Article 4, Appendix III - Phased Array Ultrasonic Examination\n\n적용 범위:\n위상배열(Phased Array) 전용 부록으로, PA 시스템 고유 요구사항 규정.\n\n주요 요구사항:\n- 초점 법칙(Focal Law) 설계 및 검증\n- 섹터 스캔(S-scan) 각도 범위 및 해상도 설정\n- 선형 스캔(Linear Scan) 인덱스 포인트 설정\n- 유효 빔 프로파일 검증 (Beam Profile Verification)\n- 개구수(Aperture) 및 소자 수 설정 기준\n- 데이터 획득 설정: 피치, 펄스 반복 주파수\n\n비고: Article 4 본문과 함께 적용 필수",
            "ASME Section VIII Div.1, Appendix 12 (PAUT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 12\n- Ultrasonic Examination of Welds (압력용기 용접부 초음파 검사)\n\n적용 범위:\n압력용기 용접부에 대한 RT 대체 UT(PAUT 포함) 적용 기준.\n\n주요 요구사항:\n- 검사 범위: 용접부 전체 체적 + 열영향부(HAZ) 포함\n- 탐촉자 선정: 용접 형상, 두께, 재질에 따라 결정\n- 보정: ASME 보정 블록 사용, 검사 두께별 감도 설정\n- 수용 기준: Table UW-53 적용\n- RT 대체 적용 시: 기술 근거 및 절차 승인 필요\n\n검사원 자격: ASNT Level II 이상",
            "ASME Section VIII Div.2, Para. 7.5.5 (PAUT - 고압용기)":
                "ASME Section VIII Division 2, Paragraph 7.5.5 - Ultrasonic Examination\n\n적용 범위:\n고압용기 용접부 PAUT 체적 검사 요구사항.\n\n주요 요구사항:\n- 검사 범위: 전체 용접 체적 + HAZ 100% 검사\n- 절차: ASME Section V Article 4 및 Appendix III 준수\n- 보정: 동일 재질·두께의 기준 블록 사용\n- 감도: 2mm SDH(Side Drilled Hole) 기준 DAC 설정\n- 수용 기준: Table 7.5.5-1 적용\n- 기록: 디지털 데이터(A-scan, S-scan) 전체 보존\n\n검사원 자격: ASNT Level II 또는 III",
            "ASME B31.1 (PAUT - 발전 배관)":
                "ASME B31.1 Power Piping - PAUT 적용 요구사항\n\n적용 범위:\n발전소 배관 용접부(증기, 급수, 블로우다운 등) PAUT 검사.\n\n주요 요구사항:\n- 절차: ASME Section V Article 4 + Appendix III 기반 서면 절차\n- 수용 기준: Table 136.4.1 (결함 유형별 기준)\n- 검사 범위: 136.5 조항 용접부 검사 요구사항 준수\n- RT 대체: 동등 검사 능력 기술 입증 시 PAUT로 RT 대체 가능\n- 자격: SNT-TC-1A Level II/III\n- 보정 블록: IIW 블록 또는 ASME 보정 블록\n- 기록: 결함 위치, 크기, 평가 결과, 검사 조건 문서화\n- 보고서: 절차, 장비, 인원, 결과 및 판정 포함",
            "ASME B31.3 (PAUT - 공정 배관)":
                "ASME B31.3 Process Piping - PAUT 적용 요구사항\n\n적용 범위:\n석유화학·정유·가스 플랜트 공정 배관 용접부 PAUT 검사.\n\n주요 요구사항:\n- 배관 등급별 검사 범위:\n  · Normal Fluid Service: 용접부의 5% 이상\n  · Category M / High Pressure: 100% 검사\n- 절차: ASME Section V Article 4 준수\n- 수용 기준: Table 341.3.2 적용\n- RT 대체: PAUT로 RT 대체 시 동등 이상 감도 입증 필요\n- 자격: ASNT SNT-TC-1A Level II 이상\n- 결과 기록: 검사 부위, 결함 지시, 판정 결과 문서화",
            "ASME Section XI (PAUT - 원자력 가동 중 검사)":
                "ASME Section XI - Rules for Inservice Inspection of Nuclear Power Plant Components\n\n적용 범위:\n원자력 발전소 핵심 기기·배관·용기의 가동 중 검사(ISI).\n\n주요 요구사항:\n- 검사 주기: IWB/IWC/IWD 조항별 10년 주기 검사 계획\n- PAUT 적용: Appendix VIII (성능 실증, PDI) 필수 통과\n- 절차 검증: Performance Demonstration Initiative (PDI) 인증\n- 적용 부위: 반응기 압력용기, 1차 배관, 노심 내부 구조물\n- 결함 크기 기준: IWB-3500 시리즈 수용 기준\n- 자격: ASNT Level III + PDI 자격 보유자 감독 하 수행\n- 기록: 전체 스캔 데이터 10년 이상 보존",
            "ASME Section I (PAUT - 보일러)":
                "ASME Section I - Power Boilers (PAUT 적용)\n\n적용 범위:\n발전용 보일러 동체, 헤더, 고온 배관 용접부 검사.\n\n주요 요구사항:\n- 적용 부위: 보일러 동체 용접부, 드럼, 노즐, 헤더 용접부\n- 절차: ASME Section V Article 4 기반\n- 수용 기준: PW-51 조항 적용\n- 두께 범위: 주로 25mm 이상 후판 용접부에 PAUT 적용\n- 보정: 검사 두께에 맞는 ASME 보정 블록 사용\n- 자격: ASNT SNT-TC-1A Level II 이상\n- 기록: 스캔 데이터 및 판정 결과 보존",
            "ASME Sec. VIII Div.2 Para. 7.5.5 (Ultrasonic Examination)": 
                "ASME Section VIII Division 2, Paragraph 7.5.5: Ultrasonic Examination\n\nUltrasonic examination shall be performed in accordance with the requirements of Article 7.5. The examination shall be conducted by qualified personnel using calibrated equipment. Acceptance criteria shall meet the requirements of Table 7.5.5-1.",
            "ASME Sec. VIII Div.2 Para. 7.5.4 (Radiographic Examination)": 
                "ASME Section VIII Division 2, Paragraph 7.5.4: Radiographic Examination\n\nRadiographic examination shall be performed in accordance with the requirements of Article 7.5. The examination shall be conducted by qualified personnel using approved techniques. Acceptance criteria shall meet the requirements of Table 7.5.4-1.",
            "ASME B31.1 PAUT 관련 코드":
                "ASME B31.1 PAUT 관련 코드\n\n- 적용 기준: ASME B31.1 Table 136.4.1 수용 기준.\n- 검사 범위: B31.1 136.5 용접부 검사 요구 사항 준수.\n- 절차 기준: ASME V Article 4 및 Appendix III에 따른 서면 시험 절차.\n- 자격: SNT-TC-1A 또는 ASNT Level II / III 자격.\n- 장비 보정: IIW 블록 또는 ASME 보정 블록을 사용한 보정 및 감도 확인.\n- 스캔 범위: 용접부 및 열영향부 전체 체적 커버리지.\n- 기록: 결함 위치, 크기, 평가 결과, 테스트 조건을 포함한 문서화.\n- 보고서: 절차 식별, 장비, 인원, 검사 결과 및 판정 포함.",
            "ASME B31.1 PAUT (Phased Array Ultrasonic Testing)": 
                "ASME B31.1 Power Piping - PAUT Requirements\n\n1. Procedure: Written procedure in accordance with ASME V Article 4 and Appendix III.\n2. Personnel: Qualified Level II or III per SNT-TC-1A, ASNT, or equivalent qualification program.\n3. Equipment: Calibrated phased array ultrasonic testing system with appropriate probes and wedge angles.\n4. Calibration: Use reference blocks such as IIW or ASME calibration blocks and perform sensitivity checks.\n5. Scanning: Ensure full volumetric coverage of the weld and heat-affected zone.\n6. Evaluation: Apply acceptance criteria in ASME B31.1 Table 136.4.1 or project-specific criteria.\n7. Records: Maintain complete inspection records, including defect sizing, location, probe data, and disposition.\n8. Reporting: Include procedure identification, equipment, personnel, and results in the test report.",
            "ASME B31.1 PAUT 필수 항목":
                "ASME B31.1 PAUT 필수 항목\n\n- 검사 절차: ASME V Article 4, Appendix III에 따른 서면 절차.\n- 자격: SNT-TC-1A 또는 ASNT 기준의 Level II/III 검사자.\n- 장비: 적절한 위상배열 초음파 시스템 및 프로브.\n- 보정: IIW 블록 또는 ASME 보정 블록을 이용한 보정 및 감도 확인.\n- 스캔: 용접부와 열영향부의 전체 체적 커버리지.\n- 평가: B31.1 Table 136.4.1 기준 또는 지정된 수용 기준.\n- 기록: 결함 위치, 크기, 평가 및 처분을 포함한 완전한 기록.\n- 보고: 절차, 장비, 인원, 결과를 포함한 보고서 작성.",
            "ASME B31.1 PAUT 합부판정 기준 (Table 136.4.1)":
                "ASME B31.1 PAUT 합부판정 기준 - Table 136.4.1\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ 불합격 지시 (Unacceptable Indications)\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 균열 (Cracks)\n- 모든 균열 지시: 크기·길이 무관하게 전부 불합격\n\n[2] 완전 용입 불량 (Incomplete Penetration)\n- 두께 기준:\n  · t < 19mm : 길이 합계 > 25mm 이상 불합격\n  · 19mm ≤ t < 57mm : 길이 합계 > t/3 이상 불합격  \n  · t ≥ 57mm : 길이 합계 > 19mm 이상 불합격\n  (임의 300mm 구간 내 누적 길이 기준)\n\n[3] 용합 불량 (Incomplete Fusion)\n- 완전 용입 불량과 동일 기준 적용\n\n[4] 내부 결함 (Internal Defects - 기공, 슬래그 등)\n- 개별 지시 길이 > 6mm 불합격\n- 임의 150mm 구간 내 지시 길이 합계 > 12mm 불합격\n- 단, 최대 개별 지시 < 3mm이고 군집 면적 < 6cm² 이하는 허용\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ 합격 조건 (Acceptable Indications)\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 위 불합격 기준에 해당하지 않는 모든 지시\n- PAUT 진폭 기반 평가: DAC(Distance Amplitude Correction) 20% 이하 지시는 기록 불요\n- 20~100% DAC: 기록 필요, 크기 측정 후 위 기준과 비교 평가\n- 100% DAC 초과: 불합격 추정, 크기 측정·평가 필수\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ PAUT 특수 고려사항\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 결함 크기 측정: -6dB drop 또는 -20dB drop 법 적용\n- 표면 결함 연결 여부 확인 필수\n- 체적 커버리지: 전체 용접 단면의 100% 스캔 데이터 확보",
            "ASME B31.3 PAUT 합부판정 기준 (Table 341.3.2)":
                "ASME B31.3 PAUT 합부판정 기준 - Table 341.3.2\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Normal Fluid Service 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 균열 (Cracks) → 전부 불합격\n\n[2] 용합 불량 / 용입 불량 (IF / IP)\n- 길이가 다음을 초과하면 불합격:\n  · t ≤ 6mm  : 2mm\n  · 6mm < t ≤ 19mm : t/3\n  · t > 19mm : 6mm\n  (임의 100mm 구간 내 누적)\n\n[3] 내부 기공·슬래그 (Porosity / Slag)\n- 개별 지시 > 3mm 불합격\n- 임의 100mm 구간 내 합계 > 6mm 불합격\n- 군집 기공: 25cm² 투영 면적 내 > 1cm² 기공 면적 불합격\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Severe Cyclic / Category M 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 용합 불량 / 용입 불량: 길이 > 0 (어떠한 지시도 불합격)\n- 기공: 개별 > 1.5mm 불합격\n- 균열: 전부 불합격\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ PAUT 진폭 평가 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 기준 감도: SDH(Side Drilled Hole) 또는 FBH(Flat Bottom Hole) DAC 설정\n- 기록 수준: DAC 20% (−14dB) 이상 모든 지시 기록\n- 평가 수준: DAC 50% (−6dB) 이상 크기 측정 및 위 기준 적용\n- DAC 100% 초과 지시: 불합격 추정, 반드시 크기 측정·평가",
            "ASME Sec. VIII Div.1 PAUT 합부판정 기준 (UW-51/App.12)":
                "ASME Section VIII Div.1 PAUT 합부판정 기준\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ UW-51 (전수 RT 대체 PAUT) 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 무조건 불합격 결함\n- 균열 (Cracks): 크기·위치 무관 전부 불합격\n- 용합 불량 (Incomplete Fusion)\n- 용입 불량 (Incomplete Penetration)\n\n[2] 내부 결함 (기공, 슬래그)\n두께(t)별 최대 허용 개별 지시 길이:\n  · t ≤ 19mm  → 최대 6mm\n  · 19mm < t ≤ 57mm → 최대 t/3\n  · t > 57mm  → 최대 19mm\n\n임의 12t 구간(단, 최대 152mm) 내 지시 길이 합계:\n  → 위 개별 기준치 이내\n\n[3] 언더컷 (Undercut)\n- 표면 언더컷 깊이 > 1mm: 불합격\n- 0.4mm 미만: 허용\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Appendix 12 (UT 대체) 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 평면형 결함 (균열, IF, IP): 전부 불합격\n- 체적형 결함 (기공, 슬래그):\n  · 개별 지시 높이 > 25% t 또는 6mm 중 작은 값: 불합격\n  · 지시 길이: UW-51 기준과 동일 적용\n- DAC 기준:\n  · 평가 수준: DAC 50%\n  · 기록 수준: DAC 20%",
            "ASME PAUT 파괴역학 합부판정 (Code Case 2235 / ECA)":
                "ASME PAUT 파괴역학 합부판정 기준\n"
                "Code Case 2235 / Engineering Critical Assessment (ECA)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 개요 (Overview)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "파괴역학(Fracture Mechanics) 기반 합부판정은 결함의 실제\n"
                "구조적 영향을 평가하여 전통적 RT/UT 기준보다 완화된 허용\n"
                "기준을 적용하는 방법. ASME Code Case 2235가 대표적.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ASME Code Case 2235 (UT/PAUT로 RT 대체)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "적용 조건:\n"
                "- 적용 코드: ASME Section VIII Div.1, Div.2 및 Section I\n"
                "- 두께 범위: 20mm ≤ t ≤ 250mm\n"
                "- 대상 용접부: 완전 용입 맞대기 용접 (Full Penetration Butt Weld)\n"
                "- 적용 재질 (P-Number):\n"
                "  P-No.1(탄소강), 3(저합금), 4(Cr-Mo), 5A/5B(고Cr-Mo),\n"
                "  6/7(마르텐사이트/페라이트 스테인리스), 8(오스테나이트 STS),\n"
                "  9A/9B(Ni합금강), 10A/10F(고강도 저합금강)\n"
                "- 용접 조인트 범주: Category A (종방향·환형 맞대기) 우선 적용\n"
                "- 표면 상태: 검사 전 기계 가공 또는 연마 (Ra ≤ 6.3 μm 권장)\n"
                "- 검사 온도: 15°C ~ 50°C (재질·절차에 따라 조정)\n\n"
                "절차 자격 인정 (Procedure Qualification):\n"
                "- ASME Section V, Article 4 Appendix III 준수 서면 절차 필수\n"
                "- 실제 검사 두께와 동일한 데모 블록(Demonstration Block)으로\n"
                "  절차 유효성 검증 (Blind Test / Mock-Up)\n"
                "- POD (Probability of Detection): 목표 결함 크기에서 ≥ 90%\n"
                "- 절차 변경(장비·프로브·각도·설정) 시 재자격 인정 필요\n\n"
                "보정 블록 기준 (Calibration Block):\n"
                "- 재질: 검사 대상과 동종 또는 음향 임피던스 동등 재질\n"
                "- 기준 반사체 (SDH, Side Drilled Hole):\n"
                "  · t ≤ 50mm  → Ø 1.5mm SDH\n"
                "  · t > 50mm  → Ø 2.0mm SDH\n"
                "- SDH 배치 깊이: t/4, t/2, 3t/4 (최소 3개 지점)\n"
                "- 보정 주기: 검사 시작 전, 종료 후, 매 4시간마다, 장비 이동 시\n"
                "- 온도 보정: 검사체와 보정 블록 온도 차 ±14°C 초과 시 재보정\n\n"
                "스캔 커버리지 (Scan Coverage):\n"
                "- 전체 용접 체적의 100% 커버리지 확보 필수\n"
                "- HAZ(열영향부) 포함 모재 측 최소 t/4 이상 구간 검사\n"
                "- 인덱스 방향 스캔 간격(Scan Increment): ≤ 1.0mm\n"
                "- S-scan 빔 각도 범위: 40°~70° (스텝 ≤ 2° 권장)\n"
                "- 모든 각도에서 -6dB 중첩(Beam Overlap) 유지\n\n"
                "신호 수준 기준 (Amplitude Criteria):\n"
                "- 기록 수준 (Recording Level)  : DAC 20% (−14 dB) 이상\n"
                "- 평가 수준 (Evaluation Level) : DAC 50% (−6 dB) 이상 → 크기 측정\n"
                "- 거부 수준 (Rejection Level)  : DAC 100% 초과 → 즉시 크기 평가\n"
                "- S/N 비: ≥ 3:1 (9.5 dB) 유지 필수\n\n"
                "결함 허용 기준 (Allowable Flaw Size):\n"
                "아래 조건을 모두 만족 시 합격\n\n"
                "  (1) 결함 높이 (a, Through-Thickness):\n"
                "      a ≤ 0.1t  (단, a ≤ 6mm)\n\n"
                "  (2) 결함 길이 (ℓ, Along Weld):\n"
                "      ℓ ≤ 6a  (최대 50mm)\n\n"
                "  (3) 표면 연결 결함 (Surface-Breaking):\n"
                "      허용 높이 기준 50% 감소 → a_allow × 0.5 적용\n\n"
                "  (4) 결함 간격 규칙 (Flaw Spacing Rule):\n"
                "      인접 결함 간격 S < max(a₁, a₂) 이면\n"
                "      두 결함을 단일 결함으로 합산하여 평가\n\n"
                "평면형 결함 (Planar Flaws - 균열, 용합불량, 용입불량):\n"
                "  위 (1)~(4) 기준 동시 적용\n"
                "  균열성 결함: 보수적 평가 필수, 재질의 K_IC, ΔK_th 확인\n\n"
                "검사원 자격 (Personnel Qualification):\n"
                "- ASNT SNT-TC-1A 또는 CP-189 기준 PAUT Level II 이상\n"
                "- 해당 장비·소프트웨어 교육 이수 및 실기 평가 기록 보유\n"
                "- 합부판정은 Level II 이상만 수행 가능\n\n"
                "보고서 요건 (Reporting Requirements):\n"
                "- 장비 식별(SN), 프로브 사양, 보정 데이터 첨부\n"
                "- A-scan, S-scan 원시 데이터(Raw Data) 파일 보존\n"
                "- 검사된 용접부 번호, 길이, 커버리지 맵 포함\n"
                "- 지시 목록: 위치(X, Y, 깊이), 높이, 길이, 판정 결과\n"
                "- 데이터 저장 형식: DICONDE 또는 제조사 전용 포맷\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 파괴역학 ECA (Engineering Critical Assessment)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "ECA 적용 기준 (BS 7910 / API 579 방법론):\n\n"
                "[1] 임계 결함 크기 계산 (Critical Flaw Size)\n"
                "    K_I = Y × σ × √(π·a)  ≤  K_IC / SF\n"
                "    - K_I : 응력 확대 계수\n"
                "    - Y   : 형상 계수 (결함 형상·위치 의존)\n"
                "    - σ   : 작용 응력 (막응력 + 굽힘응력 + 잔류응력)\n"
                "    - K_IC: 파괴인성 (재질별 시험값 또는 Charpy 변환식)\n"
                "    - SF  : 안전계수 (일반적으로 2.0~2.5)\n\n"
                "[2] 피로 균열 성장 평가\n"
                "    da/dN = C × (ΔK)^m  (Paris Law)\n"
                "    - 설계 수명 내 결함 성장 후 최종 크기 ≤ a_critical\n"
                "    - 검사 주기 결정에도 활용\n\n"
                "[3] 파손 평가 다이어그램 (FAD - Failure Assessment Diagram)\n"
                "    Kr = K_I / K_IC  (파괴비)\n"
                "    Lr = P / P_L     (소성 붕괴비)\n"
                "    → FAD 곡선 내부에 위치 시 합격\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ PAUT 결함 측정 요건 (ECA 입력 데이터)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- 결함 높이(a): TOFD 또는 -6dB drop 법, 정확도 ±1mm 이내\n"
                "- 결함 길이(ℓ): -6dB 또는 -20dB 끝점법\n"
                "- 결함 위치: 표면 연결 여부, 깊이(d), 두께 위치 확인\n"
                "- 결함 형상: 타원형(Elliptical) 가정 → a/c 비율 결정\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 적용 참고 코드\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "- ASME Code Case 2235 (최신 개정판 확인 필수)\n"
                "- API 579-1/ASME FFS-1 : Part 9 (균열형 결함 FFS)\n"
                "- BS 7910 : Guide to methods for assessing the acceptability\n"
                "  of flaws in metallic structures\n"
                "- ASME Section XI : IWB-3600 (원자력 배관 ECA)\n"
                "- 적용 시 재질의 파괴인성(K_IC) 시험값 또는\n"
                "  Charpy → K_IC 변환식 사용 근거 문서화 필수",

            "ASME Sec. VIII Div.2 PAUT 합부판정 기준 (Para. 7.5.5)":
                "ASME Section VIII Div.2 PAUT 합부판정 기준 - Paragraph 7.5.5\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Table 7.5.5-1 수용 기준\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n[1] 평면형 결함 (Planar Defects)\n- 균열, 용합 불량, 용입 불량: 크기 무관 전부 불합격\n\n[2] 체적형 결함 (Volumetric Defects)\n결함 높이(a) 및 길이(ℓ) 기준:\n  · a ≤ 3mm : ℓ ≤ 6mm 허용\n  · 3mm < a ≤ 6mm : ℓ ≤ 2a 허용\n  · a > 6mm : 불합격\n\n표면 결함 (Surface-Breaking):\n  · a ≤ 1.5mm : ℓ ≤ 6mm 허용\n  · a > 1.5mm : 불합격\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ PAUT 결함 크기 측정 요건\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 결함 높이(a): -6dB drop법 또는 TOFD(Tip Diffraction) 병행\n- 결함 길이(ℓ): -6dB drop법 또는 -20dB 끝점법\n- 위치 정확도: ±1mm 이내\n\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ 보정 블록 기준 (Reference Sensitivity)\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n- 2mm SDH(Side Drilled Hole)을 기준 반사체로 DAC 설정\n- 평가 수준: DAC 50% (-6dB)\n- 기록 수준: DAC 20% (-14dB)\n- 검사 감도: 평가 수준보다 +6dB 추가 증폭하여 검사",
            "RT 표준 항목":
                "RT (Radiographic Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 2 (또는 프로젝트 지정 기준).\n- 필름/디지털 기법: 노출 조건, IQI, 감도 및 판독 조건 준수.\n- 평가 기준: 해당 제작 코드(예: ASME B31.1, ASME VIII)의 수용 기준 적용.\n- 기록 및 보고: 노출 조건, 판독 결과, 결함 위치/길이 및 판정 포함.",
            "ASME Section V, Article 2 (RT 기본 절차)":
                "ASME Section V, Article 2 - Radiographic Examination (RT 기본 절차)\n\n적용 범위:\n방사선투과검사(RT)의 핵심 절차 요구사항을 규정하는 기본 코드.\n\n주요 요구사항:\n- 선원(Source): X선관 또는 감마선원(Ir-192, Co-60, Se-75) 선택 기준\n- 기하학적 조건: 선원-피사체 거리(SFD), 필름-피사체 거리 설정\n- IQI (Image Quality Indicator): 선형(Wire) 또는 구멍형(Hole) IQI 선택 및 배치\n- 최소 감도: 2-2T (Wire IQI) 또는 2% 구멍 IQI 기준 충족\n- 필름: ASTM E1815 Class 1 또는 2 이상\n- 현상 처리: 온도·시간 기준 준수 또는 디지털 이미징 시스템 적용\n- 판독 조건: 밝기 조도, 차폐 조건 규정\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME Section V, Article 2, Appendix D (디지털 RT)":
                "ASME Section V, Article 2, Appendix D - Digital Radiography (DR)\n\n적용 범위:\n디지털 방사선 시스템(DR/CR)을 이용한 RT 절차.\n\n주요 요구사항:\n- 디지털 검출기: DR(평판), CR(이미징 플레이트) 사용 가능\n- IQI: 전통 RT와 동일 요건, 디지털 콘트라스트 감도 별도 검증\n- 공간 해상도: Basic Spatial Resolution(BSR) 측정 및 기준 충족\n- 시스템 검증: 동일 두께 필름 RT와 동등성 입증\n- 데이터 저장: 무결성 보장 포맷(DICONDE 권장)으로 보존\n- 판독: 모니터 최소 해상도 및 밝기 기준 준수",
            "ASME B31.1, Para. 136.4 (RT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Radiographic Examination (발전 배관 RT)\n\n적용 범위:\n발전소 배관 용접부에 대한 RT 요구사항.\n\n주요 요구사항:\n- 검사 범위: Table 136.4 기준 (범주별 RT 비율 규정)\n  · P1 ~ P15 재질 및 두께별 의무 검사 비율\n- 절차: ASME Section V Article 2 준수\n- 수용 기준: Table 136.4.1 적용\n  · 균열, 융합불량, 미용융, 기공 등 유형별 기준\n- 100% RT 대상: 고온·고압 배관, 카테고리 D 이상\n- 보고: 노출 조건, IQI 확인, 판독 결과 기록\n\n검사원 자격: ASNT Level II 이상",
            "ASME B31.3, Para. 344.5 (RT - 공정 배관)":
                "ASME B31.3, Paragraph 344.5 - Radiographic Examination (공정 배관 RT)\n\n적용 범위:\n석유화학·정유·가스 플랜트 공정 배관 RT.\n\n주요 요구사항:\n- 검사 비율 (Table 341.3.2):\n  · Normal: 5% 이상 무작위 검사\n  · Severe Cyclic: 100%\n  · Category M: 100%\n  · High Pressure: 100%\n- 절차: ASME Section V Article 2 준수\n- 수용 기준: Table 341.3.2 및 Appendix A 규정\n- 선원: X-Ray 또는 Gamma-Ray 프로젝트 승인 조건 사용\n- 추가 검사: 거부 지시 발견 시 동일 용접사 전수 검사",
            "ASME Sec. VIII Div.1, UW-51/52 (RT - 압력용기)":
                "ASME Section VIII Division 1, UW-51/52 - Radiographic Examination (압력용기 RT)\n\n적용 범위:\n압력용기 동체·헤드·노즐 용접부 RT.\n\n주요 요구사항:\n- UW-11: 필수 RT 대상 용접부 결정 기준\n  · P-No.1 ~ 15 재질 및 두께 기준\n- UW-51 (전수 RT): 전체 용접 길이 100% 검사\n  · 1.0 이음효율(E=1.0) 적용 가능\n- UW-52 (부분 RT): 용접 길이의 일부 검사\n  · 0.85 이음효율(E=0.85) 적용\n- 수용 기준: UW-51(b) - 균열, 미융합, 불완전 용입 불허\n  · 기공: Table UW-51 면적 기준\n- 기록: 투과사진(필름 또는 디지털) 3년 이상 보존",
            "MT 표준 항목":
                "MT (Magnetic Particle Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 7 (또는 프로젝트 지정 기준).\n- 자분 방식: 건식/습식, 형광/비형광, 자화 방법 및 방향성 확인.\n- 평가 기준: 해당 제작 코드의 표면 결함 수용 기준 적용.\n- 기록 및 보고: 자화 조건, 사용 매질, 지시 길이/위치 및 판정 포함.",
            "ASME Section V, Article 7 (MT 기본 절차)":
                "ASME Section V, Article 7 - Magnetic Particle Examination (MT 기본 절차)\n\n적용 범위:\n자분탐상검사(MT)의 핵심 절차 요구사항.\n\n주요 요구사항:\n- 자화 방법:\n  · 연속법(Continuous Method): 자화 유지 중 자분 적용\n  · 잔류법(Residual Method): 자화 후 자분 적용 (고보자력 재질)\n- 자화 방식: 요크(Yoke), 코일, 헤드샷, 프로드(Prod) 선택\n- 자분: 건식(Dry Powder) 또는 습식(Wet Suspension), 형광·비형광\n- 자장 강도: 요크 - 4.5kgf(AC) 또는 18kgf(DC) 리프팅 파워 확인\n- 검사 방향: 최소 2방향(90° 교차) 자화 필요\n- 조명: 가시광선 MT - 최소 100fc(1000 lux), 형광 MT - 최대 2fc 암실\n- 온도: 자분 적용 표면 10~52°C 범위\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME B31.1, Para. 136.4 (MT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Magnetic Particle Examination (발전 배관 MT)\n\n적용 범위:\n발전소 배관 용접부 표면 결함 MT 검사.\n\n주요 요구사항:\n- 적용 부위: 표면 및 표면 직하 결함 검출 (최대 3mm 깊이)\n- 절차: ASME Section V Article 7 준수\n- 수용 기준: Table 136.4.1 (선형·원형 지시 기준)\n  · 선형 지시: 1.6mm 이상 불허\n  · 원형 지시: 직경 4.8mm 이상 불허\n- 자화 방법: 요크, 코일, 프로드 (프로드는 균열 민감 재질 주의)\n- 후처리: 검사 후 잔류 자분 제거 및 탈자 필요 시 실시",
            "ASME B31.3, Para. 344.3 (MT - 공정 배관)":
                "ASME B31.3, Paragraph 344.3 - Magnetic Particle Examination (공정 배관 MT)\n\n적용 범위:\n공정 배관 용접부 및 모재 표면 결함 MT.\n\n주요 요구사항:\n- 절차: ASME Section V Article 7 기반\n- 검사 범위: Table 341.3.2 기준 (배관 등급별 MT 비율)\n- 수용 기준: Table 341.3.2 또는 Appendix A 적용\n  · 선형 지시 2mm 이상, 원형 지시 5mm 이상 불허\n- 강자성체 재질(탄소강, 저합금강) 전용\n- 오스테나이트계 스테인리스강: PT 적용 (MT 불가)\n- 후열처리 후 잔류 자분 및 자장 제거 확인",
            "ASME Sec. VIII Div.1, App. 6 (MT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 6 - MT (압력용기)\n\n적용 범위:\n압력용기 용접부 및 표면 결함 MT.\n\n주요 요구사항:\n- 적용 시점: 용접 완료 후, 최종 열처리 후 검사\n- 검사 면적: 용접부 + 양쪽 열영향부 각 13mm 포함\n- 자화 방법: 요크(AC 권장), 코일, 헤드샷\n- 자분: 형광 습식 자분 권장 (감도 우수)\n- 수용 기준: Appendix 6, Para. 6-5 적용\n  · 선형 지시 1.6mm, 원형 지시 4.8mm 이상 불허\n  · 4개 이상 지시 열 배열 불허\n- 기록: 지시 위치, 크기, 자화 조건, 판정 포함",
            "PT 표준 항목":
                "PT (Penetrant Testing) 표준 항목\n\n- 적용 코드: ASME Section V Article 6 (또는 프로젝트 지정 기준).\n- 절차: 전처리, 침투 시간, 제거, 현상, 관찰 시간 조건 준수.\n- 평가 기준: 해당 제작 코드의 표면 결함 수용 기준 적용.\n- 기록 및 보고: 재료, 표면 상태, 지시 위치/크기 및 판정 포함.",
            "ASME Section V, Article 6 (PT 기본 절차)":
                "ASME Section V, Article 6 - Liquid Penetrant Examination (PT 기본 절차)\n\n적용 범위:\n침투탐상검사(PT) 핵심 절차 요구사항. 비자성체 포함 모든 금속 재질 적용 가능.\n\n주요 요구사항:\n- PT 시스템 종류:\n  · Type 1 (형광): 자외선(UV-A, 320~400nm) 조사 하 관찰\n  · Type 2 (가시광선): 가시광선 하 관찰\n- 제거 방법:\n  · Method A: 수세성(Water Washable)\n  · Method B: 후유화성(Post-emulsifiable, 지용성)\n  · Method C: 용제 제거성(Solvent Removable)\n  · Method D: 후유화성 (수용성)\n- 절차 단계: 전처리 → 침투 → 제거 → 현상 → 관찰 → 후처리\n- 침투 시간: 최소 5분 (재질·온도에 따라 조정)\n- 현상 시간: 10분 이상 (비수성 습식 현상제)\n- 온도: 표면 10~52°C 범위\n- 조명: 가시광선 PT - 최소 100fc(1000 lux)\n\n검사원 자격: ASNT SNT-TC-1A Level II 이상",
            "ASME B31.1, Para. 136.4 (PT - 발전 배관)":
                "ASME B31.1, Paragraph 136.4 - Liquid Penetrant Examination (발전 배관 PT)\n\n적용 범위:\n발전소 배관 용접부 표면 개구 결함 PT 검사. 비자성 재질(STS, 니켈합금 등) 주로 적용.\n\n주요 요구사항:\n- 절차: ASME Section V Article 6 준수\n- 검사 대상: 용접부 표면 + 열영향부 각 13mm 이상\n- 수용 기준: Table 136.4.1 (선형·원형 지시 기준)\n  · 선형 지시(장축/단축 ≥ 3): 1.6mm 이상 불허\n  · 원형 지시: 직경 4.8mm 이상 불허\n  · 4개 이상 지시 열 배열 불허\n- 적용 시점: 용접 완료 후 최소 24시간 경과 권장 (지연 균열 고려)\n- 후처리: 검사 후 침투제·현상제 완전 제거",
            "ASME B31.3, Para. 344.4 (PT - 공정 배관)":
                "ASME B31.3, Paragraph 344.4 - Liquid Penetrant Examination (공정 배관 PT)\n\n적용 범위:\n공정 배관 용접부·모재 표면 결함 PT. 오스테나이트계·비철금속 배관에 주로 사용.\n\n주요 요구사항:\n- 절차: ASME Section V Article 6 기반\n- 검사 범위: Table 341.3.2 기준 (배관 등급별 PT 비율)\n- 수용 기준: Table 341.3.2 또는 Appendix A\n  · 선형 지시 2mm 이상, 원형 지시 5mm 이상 불허\n- 적용 대상: 비자성체(STS 304/316, Inconel, 알루미늄 등)\n- 저온 PT: 10°C 미만 시 특수 저온 침투제 사용 및 별도 검증\n- 기록: 침투제 종류·로트번호, 침투 시간, 지시 위치·크기, 판정",
            "ASME Sec. VIII Div.1, App. 8 (PT - 압력용기)":
                "ASME Section VIII Division 1, Mandatory Appendix 8 - PT (압력용기)\n\n적용 범위:\n압력용기 용접부 및 표면 결함 PT.\n\n주요 요구사항:\n- 적용 시점: 용접 완료 후, 열처리 완료 후 최종 검사\n- 검사 면적: 용접부 + 양쪽 각 13mm 열영향부 포함\n- PT 방법: 형광 PT 권장 (감도 우수), 가시광선 PT 허용\n- 제거 방법: 용제 제거성(Method C) 또는 수세성(Method A) 적용\n- 수용 기준: Appendix 8, Para. 8-4\n  · 선형 지시 1.6mm, 원형 지시 4.8mm 이상 불허\n- 검사원: ASNT Level II 이상\n- 기록: 지시 위치·크기, 침투제 정보, 검사 조건 및 판정",
            "PMI 표준 항목":
                "PMI (Positive Material Identification) 표준 항목\n\n- 적용 기준: 프로젝트 재질 관리 절차 및 관련 코드 요구사항.\n- 장비: XRF/OES 등 교정된 장비 사용, 표준 시편으로 정확도 확인.\n- 판정: 규정 재질 성분 범위와 비교하여 적합/부적합 판정.\n- 기록 및 보고: 부재 식별번호, 측정값, 장비 정보, 검사자 및 판정 포함.",
            "API RP 578 (PMI - 공정 배관·압력용기)":
                "API Recommended Practice 578 - Material Verification Program for New and Existing Alloy Piping Systems\n\n적용 범위:\n공정 배관 및 압력용기 합금 재질 검증 프로그램의 핵심 기준.\n\n주요 요구사항:\n- 적용 대상: 합금강(P-No. 3 이상), 스테인리스강, 니켈합금, 이종금속 용접부\n- PMI 장비:\n  · XRF(X선 형광분석): 비파괴, 현장 적용 용이, 경원소 검출 한계\n  · OES(광학 방출 분광): 파괴적, 탄소 검출 가능\n- 검사 범위:\n  · 신규 배관: 합금 성분 재질 100% PMI 권장\n  · 기존 배관: 위험 기반 PMI(Risk-Based PMI) 프로그램 수립\n- 판정: 재질 규격(ASTM, ASME) 성분 범위와 비교\n- 기록: 부재 번호, 측정값, 장비 S/N, 검사자, 날짜, 판정\n- 불합격: 즉시 격리, 재질 확인 후 교체 또는 재시험",
            "ASME B31.3, Para. 342.2 (PMI - 공정 배관 재질 관리)":
                "ASME B31.3, Paragraph 342.2 - Material Identification (공정 배관 PMI)\n\n적용 범위:\n공정 배관 재질 식별 및 검증 요구사항.\n\n주요 요구사항:\n- 재질 확인 의무: 설계 사양과 다른 재질 혼용 방지\n- PMI 적용 대상:\n  · 합금강 배관 구성품(파이프, 피팅, 플랜지, 밸브)\n  · 이종금속 용접부\n- 장비: XRF 또는 OES 교정된 장비 사용\n- 검사 시점: 제작 중, 설치 전, 최종 검사 단계\n- 적용 코드와 병행: API RP 578 권고사항 반영\n- 기록 유지: 전체 PMI 결과를 배관 사양서(Line List)와 연계 관리\n- 재질 표식: PMI 확인 후 색상 코딩 또는 금속 태그 부착",
            "ASTM E1476 (PMI - XRF 기법)":
                "ASTM E1476 - Standard Guide for Metals Identification, Grade Verification, and Sorting (PMI XRF 기법)\n\n적용 범위:\nXRF(X선 형광분석)를 이용한 금속 재질 식별·등급 검증 가이드.\n\n주요 요구사항:\n- 장비 종류:\n  · 휴대형 XRF(pXRF): 현장 비파괴 분석 (Ni, Cr, Mo, V, Nb 등 검출)\n  · 벤치탑 XRF: 실험실용, 정밀도 높음\n- 교정: 매 측정 전 NIST 추적 가능 인증 표준 시편으로 교정\n- 측정 조건: 측정 시간, 면적, 표면 상태(스케일 제거 필요) 설정\n- 한계:\n  · 탄소(C), 황(S), 인(P): XRF로 검출 불가 → OES 병행 필요\n  · 도막·산화층: 표면 연마 후 측정\n- 재질 판정: ASTM, ASME, EN 규격 성분 데이터베이스 비교\n- 기록: 측정값, 장비 모델·S/N, 교정 결과, 검사자, 날짜",
            "ASME PCC-2 (PMI - 수리·교체 재질 검증)":
                "ASME PCC-2 - Repair of Pressure Equipment and Piping (수리·교체 PMI)\n\n적용 범위:\n압력 기기·배관 수리 및 교체 시 재질 검증 요구사항.\n\n주요 요구사항:\n- 적용 시점: 수리 전 기존 모재 확인, 수리 후 교체 재질 최종 검증\n- PMI 필수 대상:\n  · 합금강·고합금 재질 수리 부위\n  · 이종금속 용접 수리\n  · 재질 불명 부품 교체\n- 장비: XRF(주) + OES(탄소 확인 필요 시 보완)\n- 절차:\n  1) 기존 모재 재질 확인\n  2) 수리 재료 입고 검증\n  3) 용접 완료 후 용착금속 및 HAZ 인접 모재 재확인\n- 기록: 수리 전·후 PMI 결과, 장비 정보, 검사자, 판정 포함\n- 불합격 조치: 즉시 작업 중단, 재질 재확인 후 적합 재료로 교체",
            "기타": "사용자 정의 텍스트를 입력하세요.",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PAUT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 13588 (PAUT - 용접부 검사)":
                "ISO 13588:2019 - Non-destructive testing of welds - Ultrasonic testing - Use of automated phased array technology\n\n적용 범위:\n용접부 자동 위상배열 초음파검사(PAUT) 국제표준. ASME와 병행 사용.\n\n주요 요구사항:\n- 적용 두께: 6mm 이상 금속 용접부\n- 검사 레벨:\n  · Level A: 기본 감도, 일반 산업 적용\n  · Level B: 향상된 감도, 중요 구조물\n  · Level C: 최고 감도, 원자력·안전 중요 기기\n- 스캔 방식: 선형(Linear), 섹터(Sectorial), 복합 스캔\n- 보정 블록: 동일 재질, SDH 직경 = 검사 두께의 1/40 (최소 1mm)\n- 커버리지: 전체 용접 단면 체적 100% 스캔\n- 기록: 교정 기록, 스캔 데이터(A-scan/S-scan), 결함 지시 기록\n\n합부판정:\n- ISO 11666 (수용 기준 코드)과 연계 적용\n- 검사 레벨에 따른 기록·평가 수준 차등 적용",
            "ISO 11666 (PAUT 합부판정 수용 기준)":
                "ISO 11666:2018 - Non-destructive testing of welds - Ultrasonic testing - Acceptance levels\n\n적용 범위:\n초음파검사(UT/PAUT) 용접부 합부판정 수용 기준 국제표준.\n\n수용 레벨 (Acceptance Levels):\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 평면형 결함 (균열, 용합불량, 용입불량):\n  · 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준 반사체(SDH)의 100%: 불합격\n  · 길이: 10mm 이상 불합격\n\n■ Level 2 (표준)\n- 평면형 결함: 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준의 100%: 불합격\n  · 길이: 두께별 기준 (t < 15mm → 15mm, t ≥ 15mm → t)\n\n■ Level 3 (완화)\n- 평면형 결함: 모든 지시 불합격\n- 체적형 결함:\n  · 반사 진폭 ≥ 기준의 100%: 불합격\n  · 길이: 25mm 또는 t 중 큰 값 이상 불합격\n\n기록 수준: 기준 반사체 대비 −6dB 이상 모든 지시 기록\n평가 수준: 기준 반사체 대비 0dB 초과 지시 평가",
            "ISO 19285 (PAUT - 합부판정 대체 기준)":
                "ISO 19285:2017 - Non-destructive testing of welds - Phased array ultrasonic testing (PAUT) - Acceptance levels\n\n적용 범위:\nPAUT 전용 합부판정 국제표준. ISO 11666의 PAUT 특화 보완 기준.\n\n주요 내용:\n- PAUT 결함 크기 측정 방법 규정:\n  · -6dB 강하법 (Half Maximum Amplitude Method)\n  · -20dB 강하법 (끝점 탐지법)\n  · TOFD (Time of Flight Diffraction) 병행 가능\n- 결함 높이(a) 기반 합부판정:\n  · a < 0.5t 이고 a < 6mm: 길이 기준 추가 적용\n  · 표면 연결 결함: 길이 기준 강화\n- 스캔 인덱스 포인트 정확도: ±1mm 이내\n- 검사 레벨 (ISO 13588 Level A/B/C)과 연계\n\n비고: ASME Code Case 2235와 비교하여 적용 가능",
            "ISO PAUT 파괴역학 합부판정 (BS 7910 / ECA)":
                "ISO/BS 기반 PAUT 파괴역학 합부판정\n"
                "BS 7910 & ECA (Engineering Critical Assessment)\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ BS 7910 개요\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "BS 7910:2019 - Guide to methods for assessing the\n"
                "acceptability of flaws in metallic structures\n\n"
                "파괴역학 이론을 적용하여 결함의 실제 구조 안전성을\n"
                "평가하는 국제 표준 방법론. ISO 국제화 진행 중.\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ ECA 적용 레벨 (3단계)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Level 1 (간편법 - 보수적):\n"
                "  - 별도 재료 물성 없이 표준 FAD 곡선 적용\n"
                "  - 간단한 응력 상태, 보수적 안전계수\n"
                "  - 결과: 합격/불합격 여부만 판단\n\n"
                "Level 2 (표준법 - 일반 적용):\n"
                "  - 재료의 항복강도, 인장강도, 파괴인성(K_IC 또는 J_IC) 필요\n"
                "  - FAD 식: f(Lr) = (1 - 0.14·Lr²)[0.3 + 0.7·exp(-0.65·Lr⁶)]\n"
                "  - Kr = K_I / K_mat ≤ f(Lr)  → FAD 내부 시 합격\n"
                "  - Lr = σ_ref / σ_Y  (소성 붕괴 비율)\n\n"
                "Level 3 (정밀법 - J 적분 / R-곡선):\n"
                "  - J 적분 또는 CTOD 파괴인성 데이터 필요\n"
                "  - 연성 찢김(Ductile Tearing) 포함 평가\n"
                "  - 피로·크리프 복합 손상 평가 가능\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 결함 특성화 (PAUT 측정값 입력)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "PAUT → ECA 입력 파라미터:\n"
                "  - 결함 높이 (a): TOFD / -6dB drop 법\n"
                "  - 결함 길이 (2c): -6dB 또는 -20dB 끝점법\n"
                "  - 결함 깊이 (d): 표면 기준 결함 상단 위치\n"
                "  - 결함 형상: 타원형 (a/c ≤ 1 가정)\n"
                "  - 표면 연결 여부: 표면 결함 vs 매립 결함 구분\n\n"
                "결함 보수화 (Flaw Characterisation):\n"
                "  - 측정 불확도 고려: a_char = a_measured + δa (측정 불확도)\n"
                "  - 인접 결함 상호작용 규칙 (Interaction Criterion):\n"
                "    두 결함 간격 < min(a1, a2) 시 합산 처리\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 피로 균열 성장 평가 (BS 7910 Annex A)\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "  da/dN = A × ΔK^m  (Paris-Erdogan 법칙)\n"
                "  - ΔK_th 이하: 성장 없음\n"
                "  - 초기 결함 크기(a_0) → 설계 수명 후 최종 크기(a_f)\n"
                "  - a_f ≤ a_critical 조건 만족 시 합격\n"
                "  → PAUT 재검사 주기 결정에 활용\n\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "■ 관련 코드 / 표준\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "  - BS 7910:2019 (파괴역학 기반 ECA 주요 기준)\n"
                "  - API 579-1/ASME FFS-1 : Part 9 (균열형 결함)\n"
                "  - SINTAP / FITNET : 유럽 ECA 방법론\n"
                "  - ISO 15653 : 파괴인성 시험 (용접부 J/CTOD)\n"
                "  - ISO 12135 : K_IC 시험법\n"
                "  - DNV-RP-C210 : 해양 구조물 피로 ECA",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO RT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 17636-1 (RT - 필름 방사선 투과)":
                "ISO 17636-1:2013 - Non-destructive testing of welds - Radiographic testing - Part 1: X- and gamma-ray techniques with film\n\n적용 범위:\n필름 방사선투과검사(RT) 국제표준.\n\n주요 요구사항:\n- 검사 기법 등급:\n  · Class A (기본): 일반 산업 적용\n  · Class B (향상): 항공, 압력 기기, 안전 중요 구조물\n- 선원 종류: X선관, Ir-192, Se-75, Co-60, Yb-169\n- IQI 종류:\n  · Wire IQI (ISO 19232-1 기준)\n  · Step/Hole IQI (ISO 19232-2 기준)\n- 최소 감도: Class A - W13 wire, Class B - W14 wire 이상\n- 필름 종류: ISO 11699-1 기준 C3~C5 (Class A), C4~C6 (Class B)\n- 노출 기하: 선원-피사체 거리(f), 필름-피사체 거리(b) 규정\n- 암실 처리: 온도·시간 기준 준수\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 17636-2 (디지털 RT - DR/CR)":
                "ISO 17636-2:2013 - Non-destructive testing of welds - Radiographic testing - Part 2: X- and gamma-ray techniques with digital detectors\n\n적용 범위:\n디지털 방사선투과검사(DR: Digital Radiography, CR: Computed Radiography) 국제표준.\n\n주요 요구사항:\n- 검사 기법 등급: Class A / Class B (필름 RT와 동일 구분)\n- 디지털 검출기 종류:\n  · DR (Flat Panel Detector): 직접·간접 변환 방식\n  · CR (Storage Phosphor Imaging Plate): 이미징 플레이트\n- 기본 공간 해상도(BSR): IQI 감도와 함께 검증 필수\n- 콘트라스트 감도: 동일 두께 필름 RT와 동등 이상\n- 데이터 형식: DICONDE 형식 권장, 무손실 압축만 허용\n- 시스템 검증: 처음 사용 시 및 주요 설정 변경 시 검증\n- 영상 판독: 보정된 모니터(최소 2MP, ISO 14396 기준) 사용",
            "ISO 10675-1 (RT 합부판정 - 강재)":
                "ISO 10675-1:2016 - Non-destructive testing of welds - Acceptance levels for radiographic testing - Part 1: Steel, nickel, titanium and their alloys\n\n적용 범위:\n강재·니켈합금·티타늄 용접부 RT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 균열, 용합불량, 용입불량: 전부 불합격\n- 기공: 단독 기공 직경 ≤ 0.25t (최대 3mm), 군집 면적 ≤ 1%\n- 슬래그: 길이 ≤ t/3 (최대 10mm)\n\n■ Level 2 (표준 - 일반 적용)\n- 균열, 용합불량, 용입불량: 전부 불합격\n- 기공: 단독 기공 직경 ≤ 0.3t (최대 4mm), 군집 면적 ≤ 2%\n- 슬래그: 길이 ≤ 0.6t (최대 25mm)\n\n■ Level 3 (완화)\n- 균열: 불합격\n- 기공: 단독 기공 직경 ≤ 0.4t (최대 6mm), 군집 면적 ≤ 4%\n- 슬래그: 길이 ≤ t (최대 50mm)\n- 용입불량: 길이 ≤ t/4",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO MT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 17638 (MT - 용접부 자분탐상)":
                "ISO 17638:2016 - Non-destructive testing of welds - Magnetic particle testing\n\n적용 범위:\n용접부 자분탐상검사(MT) 국제표준.\n\n주요 요구사항:\n- 검사 기법:\n  · 형광 MT (UV-A 365nm 조사)\n  · 비형광(가시광선) MT\n- 자화 방법: 요크(Yoke), 헤드샷, 코일, 프로드, 전류 직접통전\n- 연속법 vs 잔류법 선택 기준 명시\n- 자장 강도 확인: 자장 지시계(Field Indicator) 또는 홀(Hall) 효과 가우스미터\n  · 연속법: 표면 접선 자장 2~6 kA/m\n- 자분 입자: 건식 또는 습식 (ISO 9934-2 기준)\n- 탐지 매질(Contrast Aid): 필요 시 흰색 반사 도막 사용\n- 조명:\n  · 형광 MT: UV-A 조도 ≥ 10 W/m², 주변 가시광 ≤ 20 lux\n  · 비형광 MT: 표면 조도 ≥ 500 lux\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 23278 (MT 합부판정)":
                "ISO 23278:2015 - Non-destructive testing of welds - Magnetic particle testing - Acceptance levels\n\n적용 범위:\n용접부 MT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 선형 지시 (Linear Indications, 길이/폭 ≥ 3:1):\n  · 길이 > 1.5mm: 불합격\n- 비선형 지시 (Non-linear):\n  · 장축 > 3mm: 불합격\n- 열 배열 지시: 3개 이상 / 구간 내 합계 > 3mm 불합격\n\n■ Level 2 (표준)\n- 선형 지시: 길이 > 3mm 불합격\n- 비선형 지시: 장축 > 5mm 불합격\n- 열 배열 지시: 4개 이상 / 구간 내 합계 > 6mm 불합격\n\n■ Level 3 (완화)\n- 선형 지시: 길이 > 6mm 불합격\n- 비선형 지시: 장축 > 8mm 불합격\n\n비고: 균열 지시는 레벨 무관 전부 불합격",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PT 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 3452-1 (PT - 침투탐상 기본 절차)":
                "ISO 3452-1:2021 - Non-destructive testing - Penetrant testing - Part 1: General principles\n\n적용 범위:\n침투탐상검사(PT) 일반 원칙 국제표준.\n\n주요 요구사항:\n- PT 계열 (System):\n  · Type I: 형광 침투제 (UV-A 하 관찰)\n  · Type II: 가시광선 침투제\n- 제거 방법:\n  · Method A: 수세성 (Water Washable)\n  · Method B: 후유화성 지용성 (Lipophilic Post-emulsifiable)\n  · Method C: 용제 제거성 (Solvent Removable)\n  · Method D: 후유화성 수용성 (Hydrophilic Post-emulsifiable)\n- 현상제 종류: 건식 분말, 수용성, 비수성 습식, 특수 형광\n- 침투 시간: 최소 5분 (재질·온도 따라 조정, ISO 3452-4 참조)\n- 현상 시간: 10분 이상 (비수성 현상제 기준)\n- 온도: 10~50°C 범위\n- 시약 계열 적합성: ISO 3452-2 기준 동일 제조사 계열 사용 원칙\n\n검사원 자격: ISO 9712 Level 2 이상",
            "ISO 23277 (PT 합부판정)":
                "ISO 23277:2015 - Non-destructive testing of welds - Penetrant testing - Acceptance levels\n\n적용 범위:\n용접부 PT 합부판정 국제표준.\n\n수용 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1 (엄격)\n- 선형 지시 (Linear, 길이/폭 ≥ 3:1):\n  · 길이 > 1.5mm: 불합격\n- 비선형 지시:\n  · 장축 > 3mm: 불합격\n- 열 배열 지시: 3개 이상 / 합계 > 3mm 불합격\n\n■ Level 2 (표준)\n- 선형 지시: 길이 > 3mm 불합격\n- 비선형 지시: 장축 > 4mm 불합격\n- 열 배열 지시: 4개 이상 / 합계 > 6mm 불합격\n\n■ Level 3 (완화)\n- 선형 지시: 길이 > 6mm 불합격\n- 비선형 지시: 장축 > 8mm 불합격\n\n비고:\n- 균열 지시: 레벨 무관 전부 불합격\n- 표면 개구 여부 확인 후 선형/비선형 분류\n- 검사 후 침투제·현상제 완전 제거 확인",

            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            # ISO PMI 관련 코드
            # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
            "ISO 15011-4 (PMI - 합금 재질 식별 가이드)":
                "ISO 15011-4 - Material identification and verification methods (합금 재질 식별)\n\n적용 범위:\n합금 금속 재질 식별 및 검증을 위한 국제 가이드라인.\n\n주요 내용:\n- 재질 식별 방법 분류:\n  · XRF (X-ray Fluorescence): 비파괴, 현장 신속 분석\n  · OES (Optical Emission Spectrometry): 파괴적, 탄소 검출 가능\n  · LIBS (Laser-Induced Breakdown Spectroscopy): 비파괴, 소형화 가능\n- 분석 원소 범위:\n  · XRF: Mg ~ U (Z=12~92), 탄소·질소·산소 검출 불가\n  · OES: C, Si, Mn, P, S, Cr, Ni, Mo 등 전원소 분석 가능\n- 교정 요건: NIST/PTB 추적 인증 표준 시편 사용\n- 측정 불확도: 주요 합금 원소 ±0.05wt% 이내 권장\n- 재질 판정: ASTM, EN, JIS, KS 규격 성분 범위와 비교\n- 기록: 측정값, 장비 정보, 교정 기록, 검사자, 날짜 포함",
            "ISO 9712 (NDT 검사원 자격 인증)":
                "ISO 9712:2021 - Non-destructive testing - Qualification and certification of NDT personnel\n\n적용 범위:\n비파괴검사(NDT) 전 종목 검사원 자격 부여 및 인증 국제표준.\n적용 종목: PT, MT, RT, UT(PAUT 포함), ET(와전류), VT, ST, LT 등\n\n자격 레벨:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n■ Level 1\n- 지정된 NDT 지시서에 따라 검사 수행\n- 검사 결과 기록 (합부판정 권한 없음)\n- 자격 요건: 산업 교육 + 현장 경험 + 필기·실기 시험\n\n■ Level 2 (현장 주력 자격)\n- 검사 절차 설정 및 수행\n- 합부판정 (해당 코드 기준 적용)\n- Level 1 지도·감독\n- 자격 요건: Level 1 경험 + 교육 시간 + 시험 합격\n\n■ Level 3\n- NDT 절차·기술 개발 및 승인\n- 합부판정 기준 해석\n- Level 1/2 자격 시험 감독 및 인증\n- 자격 요건: 학력 + 광범위한 실무 경험 + 종합 시험\n\n인증 유효기간: 5년 (중간 재확인 + 갱신 시험)\n인증 기관: 각국 ISO 9712 인정 인증 기관 (한국: KSNT 등)",
            "기타": "사용자 정의 텍스트를 입력하세요."
        }
        
        # 상단 저장 영역 (최우선 표시)
        top_frame = tk.Frame(root, bg="#e8f4fd", relief=tk.RIDGE, bd=1)
        top_frame.pack(fill=tk.X, padx=10, pady=(10, 2))
        
        tk.Label(top_frame, text="문서 제목:", font=("Arial", 10, "bold"), bg="#e8f4fd").pack(side=tk.LEFT, padx=(10, 5), pady=6)
        self.new_title_entry = tk.Entry(top_frame, width=50, font=("Arial", 10))
        self.new_title_entry.pack(side=tk.LEFT, padx=5, pady=6)
        self.new_title_entry.insert(0, "비파괴 검사 절차서")
        
        tk.Button(top_frame, text="💾  Word 문서 생성 / 저장", command=self.generate_document,
                  bg="#2196F3", fg="white", padx=15, pady=5,
                  font=("Arial", 10, "bold"), relief=tk.RAISED).pack(side=tk.LEFT, padx=15, pady=6)
        
        # 하단 버튼 영역
        button_frame = tk.Frame(root)
        button_frame.pack(fill=tk.X, padx=10, pady=(2, 10))
        
        tk.Button(button_frame, text="Word 파일 로드", command=self.load_document, bg="lightblue", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="사진 추가", command=self.add_images, bg="lightgreen", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="표준 추가", command=self.add_standard, bg="orange", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="텍스트 편집", command=self.edit_selected_text, bg="#d6eaff", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="선택 삭제", command=self.delete_selected_item, bg="#ffb3b3", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="전체 초기화", command=self.clear_all, bg="lightcoral", padx=10, pady=5).pack(side=tk.LEFT, padx=5)
        
        self.status_label = tk.Label(root, text="로드된 문서: 없음 | 포함된 사진: 0개", bg="lightyellow", pady=5)
        self.status_label.pack(fill=tk.X, padx=10)
        
        # 가운데 내용 영역
        main_frame = tk.Frame(root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        left_frame = tk.Frame(main_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        right_frame = tk.Frame(main_frame)
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=False, padx=(5, 0))
        
        # 트리뷰
        tree_frame = tk.Frame(left_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        tree_scrollbar = tk.Scrollbar(tree_frame)
        tree_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree = ttk.Treeview(tree_frame, yscrollcommand=tree_scrollbar.set, height=35)
        self.tree.pack(fill=tk.BOTH, expand=True)
        tree_scrollbar.config(command=self.tree.yview)
        
        self.tree['columns'] = ('타입', '내용')
        self.tree.column('#0', width=50, anchor='w')
        self.tree.column('타입', width=80, anchor='w')
        self.tree.column('내용', width=950, anchor='w')
        self.tree.heading('#0', text='번호')
        self.tree.heading('타입', text='타입')
        self.tree.heading('내용', text='내용')
        
        self.tree.bind('<Double-1>', self.on_tree_double_click)
        self.tree.bind('<Button-1>', self.on_tree_click)
        self.tree.bind('<Button-3>', self.on_tree_right_click)  # 우클릭 바인드
        
        # 우측 내용 및 이미지 영역
        info_label = tk.Label(right_frame, text="로드된 문서 내용", font=("Arial", 11, "bold"))
        info_label.pack(anchor='w')
        
        content_frame = tk.Frame(right_frame)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        content_scrollbar = tk.Scrollbar(content_frame)
        content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.content_text = tk.Text(content_frame, width=60, wrap=tk.WORD, yscrollcommand=content_scrollbar.set)
        self.content_text.pack(fill=tk.BOTH, expand=True)
        content_scrollbar.config(command=self.content_text.yview)
        self.content_text.config(state=tk.DISABLED)
        self.content_text.bind('<Button-1>', self._on_content_text_click)
        self.content_text.bind('<Motion>', self._on_content_text_motion)
        
        # 이미지 미리보기 영역
        image_preview_label = tk.Label(right_frame, text="이미지 미리보기", font=("Arial", 11, "bold"))
        image_preview_label.pack(anchor='w', pady=(10, 0))
        
        self.image_canvas = tk.Canvas(right_frame, height=220)
        self.image_canvas.pack(fill=tk.X, pady=5)
        self.image_frame = tk.Frame(self.image_canvas)
        self.image_canvas.create_window((0, 0), window=self.image_frame, anchor='nw')
        
        self.image_scrollbar = tk.Scrollbar(right_frame, orient=tk.HORIZONTAL, command=self.image_canvas.xview)
        self.image_scrollbar.pack(fill=tk.X)
        self.image_canvas.configure(xscrollcommand=self.image_scrollbar.set)
        self.image_frame.bind('<Configure>', lambda e: self.image_canvas.configure(scrollregion=self.image_canvas.bbox('all')))
        self.image_thumbnails = []
        self.content_images = []
        
        self.update_info_text()
    
    def on_tree_double_click(self, event):
        item = self.tree.identify_row(event.y)
        if not item:
            selected = self.tree.selection()
            if not selected:
                return
            item = selected[0]
        else:
            self.tree.selection_set(item)
            self.tree.focus(item)

        self.edit_selected_text(item)

    def on_tree_click(self, event):
        item = self.tree.identify_row(event.y)
        if not item:
            selected = self.tree.selection()
            if not selected:
                return
            item = selected[0]
        else:
            self.tree.selection_set(item)
            self.tree.focus(item)

        try:
            self._highlight_content_text_item(int(item))
        except Exception:
            pass

        values = self.tree.item(item, 'values')
        if not values:
            return
        if values[0] == 'image':
            content_index = int(item)
            if 0 <= content_index < len(self.content):
                self.replace_image_dialog(content_index)
        elif values[0] == 'table':
            content_index = int(item)
            if 0 <= content_index < len(self.content):
                self.view_table_dialog(content_index)

    def _highlight_content_text_item(self, idx):
        """트리 선택 항목을 content_text에서 하이라이트 및 스크롤"""
        tag = f"item_{idx}"
        try:
            ranges = self.content_text.tag_ranges(tag)
            if ranges:
                self.content_text.tag_remove("_hl_sel", "1.0", tk.END)
                self.content_text.tag_add("_hl_sel", ranges[0], ranges[1])
                self.content_text.tag_config("_hl_sel", background="#fff3cd")
                self.content_text.see(ranges[0])
        except Exception:
            pass

    def _on_content_text_click(self, event):
        """요약 텍스트 클릭 → 트리뷰 해당 항목 선택 및 스크롤"""
        try:
            pos = self.content_text.index(f"@{event.x},{event.y}")
            tags = self.content_text.tag_names(pos)
            for tag in tags:
                if tag.startswith("item_"):
                    item_idx = int(tag[5:])
                    iid = str(item_idx)
                    if self.tree.exists(iid):
                        self._highlight_content_text_item(item_idx)
                        self.tree.selection_set(iid)
                        self.tree.focus(iid)
                        self.tree.see(iid)
                    break
        except Exception:
            pass

    def _on_content_text_motion(self, event):
        """마우스 오버 시 커서 hand2 전환"""
        try:
            pos = self.content_text.index(f"@{event.x},{event.y}")
            tags = self.content_text.tag_names(pos)
            has_item = any(t.startswith("item_") for t in tags)
            self.content_text.config(cursor="hand2" if has_item else "")
        except Exception:
            pass

    def get_insert_index(self, mode='after'):
        selected = self.tree.selection()
        if selected:
            base_idx = int(selected[0])
            idx = base_idx if mode == 'before' else base_idx + 1
            return min(max(idx, 0), len(self.content))
        return len(self.content)

    def edit_selected_text(self, item_id=None):
        item = item_id
        if item is None:
            selected = self.tree.selection()
            if not selected:
                messagebox.showwarning("선택 필요", "편집할 텍스트 항목을 선택하세요.")
                return
            item = selected[0]

        values = self.tree.item(item, 'values')
        if not values or values[0] != 'text':
            messagebox.showwarning("선택 오류", "텍스트 항목만 편집할 수 있습니다.")
            return

        content_index = int(item)
        if 0 <= content_index < len(self.content):
            current_text = self.content[content_index].get('text', '')
            self.edit_text_dialog(content_index, current_text)
    
    def edit_text_dialog(self, index, current_text):
        dialog = tk.Toplevel(self.root)
        dialog.title("텍스트 편집")
        dialog.geometry("600x400")
        
        def save_text():
            new_text = text_area.get("1.0", tk.END).rstrip("\n")
            if 0 <= index < len(self.content) and self.content[index].get('type') == 'text':
                self.content[index]['text'] = new_text
                self.refresh_content()
                # 트리뷰 선택 복원
                iid = str(index)
                if self.tree.exists(iid):
                    self.tree.selection_set(iid)
                    self.tree.focus(iid)
                    self.tree.see(iid)
                # 오른쪽 요약에서 수정된 항목 하이라이트 & 스크롤
                self._highlight_content_text_item(index)
                dialog.destroy()
        
        # 상단 툴바 (저장 버튼)
        toolbar = tk.Frame(dialog, bg="#e8f4fd", relief=tk.RIDGE, bd=1)
        toolbar.pack(fill=tk.X, padx=0, pady=(0, 5))
        
        tk.Button(toolbar, text="💾  저장", command=save_text,
                  bg="#2196F3", fg="white", padx=15, pady=4,
                  font=("Arial", 10, "bold"), relief=tk.RAISED).pack(side=tk.LEFT, padx=10, pady=5)
        tk.Button(toolbar, text="✕  닫기", command=dialog.destroy,
                  bg="#f44336", fg="white", padx=12, pady=4,
                  font=("Arial", 10), relief=tk.RAISED).pack(side=tk.LEFT, padx=5, pady=5)
        tk.Label(toolbar, text=f"항목 #{index + 1} 편집 중",
                 font=("Arial", 9), bg="#e8f4fd", fg="#555").pack(side=tk.LEFT, padx=15)
        
        text_area = tk.Text(dialog, wrap=tk.WORD)
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        text_area.insert(tk.END, current_text)
        text_area.focus_set()
    
    def replace_image_dialog(self, index):
        file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if file_path:
            if not (0 <= index < len(self.content)) or self.content[index].get('type') != 'image':
                return
            old_path = self.content[index]['path']
            # 새 이미지 복사
            import shutil
            temp_dir = os.path.dirname(old_path)
            new_path = os.path.join(temp_dir, f"replaced_{os.path.basename(file_path)}")
            shutil.copy(file_path, new_path)
            self.content[index]['path'] = new_path
            self.update_tree_view()
            self.update_image_preview()
            self.update_info_text()
    
    def load_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not file_path:
            return
        
        try:
            self.content = load_existing_doc(file_path)
            self.source_file = file_path  # 바닥글/헤더 이미지 보존을 위해 원본 경로 저장
            self.refresh_content()
            self.status_label.config(text=f"로드된 문서: {os.path.basename(file_path)} | 포함된 사진: {len(self.image_paths)}개")
            messagebox.showinfo("완료", f"문서를 로드했습니다.\n내용: {len(self.content)}개\n사진: {len(self.image_paths)}개")
        except Exception as e:
            messagebox.showerror("오류", f"문서 로드 실패: {str(e)}")
    
    def update_tree_view(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        for idx, item in enumerate(self.content):
            item_type = item.get('type')
            area = item.get('area', 'body')
            area_label = '' if area == 'body' else f"[{ '머릿글' if area == 'header' else '바닥글' }] "
            if item_type == 'text':
                text = item.get('text', '').strip().replace('\n', ' ')[:120]
                style = item.get('style', 'Normal')
                if text:
                    self.tree.insert('', tk.END, iid=str(idx), text=str(idx + 1), values=('text', f"{area_label}{style}: {text}"))
            elif item_type == 'image':
                path = item.get('path', '')
                self.tree.insert('', tk.END, iid=str(idx), text=str(idx + 1), values=('image', f"{area_label}{os.path.basename(path)}"))
            elif item_type == 'table':
                self.tree.insert('', tk.END, iid=str(idx), text=str(idx + 1), values=('table', f"{area_label}표"))
    
    def update_info_text(self):
        self.content_text.config(state=tk.NORMAL)
        self.content_text.delete("1.0", tk.END)
        self.content_images.clear()
        info_content = f"""
로드된 내용 요약:
- 총 요소 수: {len(self.content)}개
- 텍스트 문단: {len(self.paragraphs)}개
- 포함된 사진: {len(self.image_paths)}개
"""
        self.content_text.insert(tk.END, info_content)
        self.content_text.insert(tk.END, "\n---\n\n")
        
        text_counter = 1
        for idx, item in enumerate(self.content):
            area = item.get('area', 'body')
            area_label = '' if area == 'body' else f"[{ '머릿글' if area == 'header' else '바닥글' }]\n"
            tag = f"item_{idx}"
            start_pos = self.content_text.index(tk.END)
            if item['type'] == 'text':
                text = item.get('text', '').replace('\n', '\n')
                style = item.get('style', 'Normal')
                if text:
                    self.content_text.insert(tk.END, f"{text_counter}. {area_label}[{style}]\n{text}\n\n")
                    text_counter += 1
            elif item['type'] == 'image':
                path = item['path']
                if os.path.exists(path):
                    try:
                        img = Image.open(path)
                        img.thumbnail((280, 280))
                        photo = ImageTk.PhotoImage(img)
                        self.content_images.append(photo)
                        self.content_text.image_create(tk.END, image=photo)
                        self.content_text.insert(tk.END, "\n\n")
                    except Exception as e:
                        self.content_text.insert(tk.END, f"[이미지 로드 실패: {e}]\n\n")
            elif item['type'] == 'table':
                table_data = item['data']
                table_text = f"{area_label}[표]\n" + "\n".join("\t".join(row) for row in table_data) + "\n\n"
                self.content_text.insert(tk.END, table_text)
            end_pos = self.content_text.index(tk.END)
            if self.content_text.compare(start_pos, '<', end_pos):
                self.content_text.tag_add(tag, start_pos, end_pos)
        
        self.content_text.config(state=tk.DISABLED)
    
    def update_image_preview(self):
        for widget in self.image_frame.winfo_children():
            widget.destroy()
        self.image_thumbnails.clear()

        if not self.image_paths:
            placeholder = tk.Label(self.image_frame, text="로드된 이미지가 없습니다.", bg="white", width=80, height=10)
            placeholder.pack(padx=10, pady=10)
            return

        for idx, path in enumerate(self.image_paths, 1):
            if not os.path.exists(path):
                continue
            try:
                img = Image.open(path)
                img.thumbnail((180, 180))
                photo = ImageTk.PhotoImage(img)
                self.image_thumbnails.append(photo)
                frame = tk.Frame(self.image_frame, bd=1, relief=tk.RIDGE)
                frame.pack(side=tk.LEFT, padx=5, pady=5)
                label = tk.Label(frame, image=photo)
                label.pack()
                caption = tk.Label(frame, text=os.path.basename(path), wraplength=180)
                caption.pack()
            except Exception as e:
                print(f"이미지 미리보기 오류: {e}")
    
    def refresh_content(self):
        self.paragraphs = [item for item in self.content if item.get('type') == 'text']
        self.image_paths = [item.get('path', '') for item in self.content if item.get('type') == 'image']
        self.update_tree_view()
        self.update_image_preview()
        self.update_info_text()
    
    def add_standard(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("표준 절차 추가")
        dialog.geometry("640x380")
        
        tk.Label(dialog, text="추가할 표준을 선택하세요:").pack(pady=10)

        selected = self.tree.selection()
        if selected:
            selected_idx = int(selected[0])
            selected_values = self.tree.item(selected[0], 'values')
            selected_desc = selected_values[1] if selected_values and len(selected_values) > 1 else '선택 항목'
            selected_label_text = f"현재 선택: {selected_idx + 1}번 ({selected_desc[:60]})"
        else:
            selected_label_text = "현재 선택: 없음 (맨 끝에 추가됨)"

        tk.Label(dialog, text=selected_label_text, fg="blue").pack(pady=(0, 5))

        position_frame = tk.Frame(dialog)
        position_frame.pack(pady=(0, 8))
        tk.Label(position_frame, text="삽입 위치:").pack(side=tk.LEFT, padx=(0, 8))
        position_var = tk.StringVar(value='after')
        tk.Radiobutton(position_frame, text="선택 위", variable=position_var, value='before').pack(side=tk.LEFT)
        tk.Radiobutton(position_frame, text="선택 아래", variable=position_var, value='after').pack(side=tk.LEFT, padx=(8, 0))

        quick_frame = tk.Frame(dialog)
        quick_frame.pack(pady=5)

        standards = self.standards

        def add_standard_content(selected_key):
            if selected_key not in standards:
                messagebox.showwarning("오류", "선택한 표준이 없습니다.")
                return
            insert_index = self.get_insert_index(position_var.get())
            self.content.insert(insert_index, {'type': 'text', 'text': standards[selected_key], 'style': 'Normal', 'area': 'body'})
            self.refresh_content()
            if str(insert_index) in self.tree.get_children():
                self.tree.selection_set(str(insert_index))
                self.tree.focus(str(insert_index))
            self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
            messagebox.showinfo("완료", f"'{selected_key}' 내용이 {insert_index + 1}번째 위치에 추가되었습니다.")
            # 추가된 항목을 자동으로 편집 모드로 열기
            self.edit_text_dialog(insert_index, standards[selected_key])

        def open_standard_picker(title, keywords):
            picker = tk.Toplevel(self.root)
            picker.title(f"{title} 항목 선택")
            picker.geometry("750x520")

            tk.Label(picker, text=f"{title} 관련 항목을 선택하세요",
                     font=("Arial", 10, "bold")).pack(pady=(8, 4))

            def get_keys_by_source(source):
                """source: 'ASME' 또는 'ISO'"""
                return [
                    key for key in standards.keys()
                    if key != "기타"
                    and any(word.lower() in key.lower() for word in keywords)
                    and (source.upper() in key.upper())
                ]

            # ── 탭 구성 ──
            notebook = ttk.Notebook(picker)
            notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=4)

            tab_data = {}  # {'ASME': {'frame':..,'combo':..,'var':..,'preview':..}, 'ISO': {...}}

            def build_tab(source, bg_color):
                keys = get_keys_by_source(source)
                frame = tk.Frame(notebook, bg=bg_color)
                notebook.add(frame, text=f"  {source}  ")

                if not keys:
                    tk.Label(frame, text=f"{source} 관련 항목이 없습니다.",
                             fg="gray", bg=bg_color).pack(pady=30)
                    tab_data[source] = {'keys': [], 'var': None, 'combo': None, 'preview': None}
                    return

                var = tk.StringVar(value=keys[0])
                combo = ttk.Combobox(frame, values=keys, textvariable=var,
                                     state="readonly", width=85)
                combo.pack(pady=(8, 4), padx=8)

                preview = tk.Text(frame, wrap=tk.WORD, height=14, bg="#fafafa")
                preview.pack(fill=tk.BOTH, expand=True, padx=8, pady=(0, 6))

                def refresh(*args):
                    key = var.get()
                    preview.config(state=tk.NORMAL)
                    preview.delete("1.0", tk.END)
                    preview.insert("1.0", standards.get(key, ""))
                    preview.config(state=tk.DISABLED)

                var.trace_add('write', refresh)
                refresh()

                tab_data[source] = {'keys': keys, 'var': var, 'combo': combo, 'preview': preview}

            build_tab("ASME", "#fffde7")
            build_tab("ISO",  "#e8f5e9")

            # ── 공통 하단 버튼 ──
            btn_frame = tk.Frame(picker)
            btn_frame.pack(fill=tk.X, pady=6)

            def get_current_selection():
                """현재 활성 탭의 선택 key 반환"""
                tab_name = notebook.tab(notebook.select(), "text").strip()
                src = tab_name  # "ASME" or "ISO"
                data = tab_data.get(src)
                if data and data['var']:
                    return data['var'].get()
                return None

            def add_selected_from_picker():
                key = get_current_selection()
                if not key:
                    messagebox.showwarning("선택 없음", "추가할 항목을 선택하세요.", parent=picker)
                    return
                add_standard_content(key)
                picker.destroy()
                dialog.destroy()

            def delete_selected_standard():
                key = get_current_selection()
                if not key:
                    return
                if not messagebox.askyesno("삭제 확인",
                        f"'{key}' 항목을 목록에서 삭제하시겠습니까?", parent=picker):
                    return
                del standards[key]
                # 탭 갱신
                tab_name = notebook.tab(notebook.select(), "text").strip()
                data = tab_data.get(tab_name)
                if data:
                    new_keys = get_keys_by_source(tab_name)
                    data['keys'] = new_keys
                    if data['combo'] and new_keys:
                        data['combo']['values'] = new_keys
                        data['var'].set(new_keys[0])
                    elif not new_keys:
                        picker.destroy()
                        messagebox.showinfo("알림", "남은 항목이 없습니다.")

            tk.Button(btn_frame, text="✔  추가", command=add_selected_from_picker,
                      bg="#4CAF50", fg="white", padx=14, pady=4,
                      font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=10)
            tk.Button(btn_frame, text="🗑  삭제", command=delete_selected_standard,
                      bg="#ffb3b3", fg="red", padx=12, pady=4,
                      font=("Arial", 9, "bold")).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="닫기", command=picker.destroy,
                      padx=12, pady=4).pack(side=tk.RIGHT, padx=10)

        tk.Button(quick_frame, text="PAUT", width=10, bg="#f7d794", command=lambda: open_standard_picker("PAUT", ["PAUT", "Phased Array", "13588", "11666", "19285"]))
        tk.Button(quick_frame, text="RT", width=10, bg="#c7ecee", command=lambda: open_standard_picker("RT", ["RT", "Radiographic", "17636", "10675"]))
        tk.Button(quick_frame, text="MT", width=10, bg="#dff9fb", command=lambda: open_standard_picker("MT", ["MT", "Magnetic Particle", "17638", "23278"]))
        tk.Button(quick_frame, text="PT", width=10, bg="#badc58", command=lambda: open_standard_picker("PT", ["PT", "Penetrant", "3452", "23277"]))
        tk.Button(quick_frame, text="PMI", width=10, bg="#f9ca24", command=lambda: open_standard_picker("PMI", ["PMI", "Material Identification", "15011", "E1476", "PCC-2"]))

        for child in quick_frame.winfo_children():
            child.pack(side=tk.LEFT, padx=5)

        tk.Label(dialog, text="버튼을 눌러 해당 표준 내용을 선택해 추가하세요.", fg="gray").pack(pady=(12, 6))
    
    def insert_standard(self, text_area, selected):

        if selected and selected in self.standards:

            if selected == "기타":

                custom_text = simpledialog.askstring("입력", "텍스트를 입력하세요:")

                if custom_text:

                    text_area.insert(tk.INSERT, custom_text)

            else:

                text_area.insert(tk.INSERT, self.standards[selected])
    
    def view_table_dialog(self, index):

        if 0 <= index < len(self.content) and self.content[index].get('type') == 'table':

            data = self.content[index]['data']

            dialog = tk.Toplevel(self.root)

            dialog.title("표 보기")

            dialog.geometry("600x400")

            tree = ttk.Treeview(dialog)

            tree.pack(fill=tk.BOTH, expand=True)

            if data:

                num_cols = len(data[0])

                tree['columns'] = [f'col{i}' for i in range(num_cols)]

                for i in range(num_cols):

                    tree.column(f'col{i}', width=100, anchor='w')

                    tree.heading(f'col{i}', text=f'열 {i+1}')

                for row in data:

                    tree.insert('', tk.END, values=row)
    
    def add_images(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif *.bmp")])
        if file_paths:
            insert_index = self.get_insert_index()
            for path in file_paths:
                self.content.insert(insert_index, {'type': 'image', 'path': path, 'area': 'body'})
                insert_index += 1
            self.refresh_content()
            first_new_index = insert_index - len(file_paths)
            if str(first_new_index) in self.tree.get_children():
                self.tree.selection_set(str(first_new_index))
                self.tree.focus(str(first_new_index))
            self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
            messagebox.showinfo("완료", f"{len(file_paths)}개의 사진이 {first_new_index + 1}번째 위치부터 추가되었습니다.")

    def delete_selected_item(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("선택 필요", "삭제할 항목을 먼저 선택하세요.")
            return

        content_index = int(selected[0])
        if not (0 <= content_index < len(self.content)):
            messagebox.showerror("오류", "유효하지 않은 선택입니다.")
            return

        item = self.content[content_index]
        item_type = item.get('type', 'unknown')

        if not messagebox.askyesno("확인", f"선택한 {item_type} 항목을 삭제하시겠습니까?"):
            return

        del self.content[content_index]
        self.refresh_content()
        self.status_label.config(text=f"로드된 문서: 변경됨 | 포함된 사진: {len(self.image_paths)}개")
    
    def on_tree_right_click(self, event):
        """트리 항목 우클릭 메뉴"""
        item = self.tree.identify_row(event.y)
        if not item:
            return
        
        # 선택 항목 설정
        self.tree.selection_set(item)
        self.tree.focus(item)
        
        try:
            content_index = int(item)
        except (ValueError, IndexError):
            return
        
        if not (0 <= content_index < len(self.content)):
            return
        
        # 컨텍스트 메뉴 생성
        context_menu = tk.Menu(self.root, tearoff=False)
        
        item_type = self.content[content_index].get('type', 'unknown')
        
        context_menu.add_command(
            label="편집", 
            command=lambda: self.edit_selected_text(item)
        )
        context_menu.add_separator()
        context_menu.add_command(
            label="삭제",
            command=self.delete_selected_item,
            foreground="red"
        )
        
        # 메뉴 표시
        try:
            context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            context_menu.grab_release()
    
    def clear_all(self):
        if messagebox.askyesno("확인", "모든 내용을 초기화하시겠습니까?"):
            self.content = []
            self.refresh_content()
            self.status_label.config(text="로드된 문서: 없음 | 포함된 사진: 0개")
    
    def generate_document(self):
        if not self.content:
            messagebox.showerror("오류", "로드된 내용이 없습니다.")
            return
        title = self.new_title_entry.get().strip() or "비파괴 검사 절차서"

        # ── 원본 파일을 템플릿으로 사용 (헤더/바닥글 이미지 완전 보존) ──
        if self.source_file and os.path.exists(self.source_file):
            doc = Document(self.source_file)
            # 본문 내용만 삭제 (sectPr·헤더·바닥글 참조는 유지)
            body = doc.element.body
            sectPr = body.find(qn('w:sectPr'))
            for el in list(body):
                if el is not sectPr:
                    body.remove(el)
            # sectPr가 없으면 다시 추가
            if sectPr is not None and body.find(qn('w:sectPr')) is None:
                body.append(sectPr)
        else:
            doc = Document()
            # 헤더 구성 (원본 없을 때)
            try:
                header = doc.sections[0].header
                clear_story_container(header)
                for item in self.content:
                    if item.get('area') != 'header':
                        continue
                    if item['type'] == 'text':
                        header.add_paragraph(item.get('text', ''))
                    elif item['type'] == 'table':
                        add_bordered_table(header, item.get('data', []))
            except:
                pass
            # 바닥글 구성 (원본 없을 때 - 이미지 포함)
            try:
                footer = doc.sections[0].footer
                clear_story_container(footer)
                for item in self.content:
                    if item.get('area') != 'footer':
                        continue
                    if item['type'] == 'text':
                        footer.add_paragraph(item.get('text', ''))
                    elif item['type'] == 'image':
                        path = item.get('path', '')
                        if os.path.exists(path):
                            try:
                                p = footer.add_paragraph()
                                p.add_run().add_picture(path, width=Inches(5))
                            except:
                                pass
                    elif item['type'] == 'table':
                        add_bordered_table(footer, item.get('data', []))
            except:
                pass

        # ── 본문 내용 작성 ──
        inserted_elements = set()  # 동일 element 중복 삽입 방지

        for item in self.content:
            if item.get('area') in ('header', 'footer'):
                continue
            elif item['type'] == 'text':
                text = item.get('text', '').strip()
                if not text:
                    continue
                style = item.get('style', 'Normal')
                if 'Heading' in style:
                    try:
                        level = int(style.split()[-1])
                    except:
                        level = 1
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text)
            elif item['type'] == 'image':
                img_elem = item.get('element')
                if self.source_file and img_elem is not None:
                    elem_id = id(img_elem)
                    if elem_id not in inserted_elements:
                        inserted_elements.add(elem_id)
                        # 원본 단락 element deepcopy → 이미지 관계(rId) 완전 보존
                        para_copy = copy.deepcopy(img_elem)
                        _body = doc.element.body
                        _sect = _body.find(qn('w:sectPr'))
                        if _sect is not None:
                            _body.insert(list(_body).index(_sect), para_copy)
                        else:
                            _body.append(para_copy)
                else:
                    path = item.get('path', '')
                    if os.path.exists(path):
                        try:
                            doc.add_picture(path, width=Inches(5))
                            doc.add_paragraph()
                        except:
                            pass
            elif item['type'] == 'table':
                data = item.get('data', [])
                tbl_elem = item.get('element')
                if tbl_elem is not None:
                    # 원본 XML 깊은 복사 → 병합 셀, 서식, 테두리 완전 보존
                    tbl_copy = copy.deepcopy(tbl_elem)
                    body = doc.element.body
                    sect_pr = body.find(qn('w:sectPr'))
                    if sect_pr is not None:
                        body.insert(list(body).index(sect_pr), tbl_copy)
                    else:
                        body.append(tbl_copy)
                    doc.add_paragraph()
                elif data:
                    add_bordered_table(doc, data)
                    doc.add_paragraph()

        output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_file:
            doc.save(output_file)
            messagebox.showinfo("완료", f"문서가 저장되었습니다:\n{output_file}")
    
    def load_window_geometry(self):
        """저장된 창 크기 로드"""
        try:
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    geometry = config.get('window_geometry', '1400x850')
                    return geometry
        except:
            pass
        return '1400x850'
    
    def save_window_geometry(self):
        """현재 창 크기 저장"""
        try:
            geometry = self.root.geometry()
            config = {}
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
            config['window_geometry'] = geometry
            with open(self.CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except:
            pass
    
    def on_window_configure(self, event):
        """창 크기 변경 시 저장"""
        if event.widget == self.root:
            self.save_window_geometry()

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = NDTProcedureApp(root)
        root.mainloop()
    except KeyboardInterrupt:
        pass
